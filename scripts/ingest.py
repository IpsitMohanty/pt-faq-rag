import os
from docx import Document

from langchain_chroma import Chroma
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_core.documents import Document as LCDocument


DOCX_PATH = os.getenv("DOCX_PATH", "backend/data/pt_faq.docx")
PERSIST_DIR = os.getenv("CHROMA_PERSIST_DIR", "backend/chroma_store")
COLLECTION = os.getenv("CHROMA_COLLECTION", "pt_faq")
EMBED_MODEL = os.getenv("EMBED_MODEL", "sentence-transformers/all-MiniLM-L6-v2")


def load_paragraphs(path):
    doc = Document(path)
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()]


def chunk(paragraphs, max_chars=900):
    chunks = []
    buf = ""

    for p in paragraphs:
        if len(buf) + len(p) < max_chars:
            buf += "\n\n" + p
        else:
            chunks.append(buf.strip())
            buf = p

    if buf.strip():
        chunks.append(buf.strip())

    return chunks


def main():
    if not os.path.exists(DOCX_PATH):
        raise FileNotFoundError(f"DOCX not found: {DOCX_PATH}")

    paras = load_paragraphs(DOCX_PATH)
    chunks = chunk(paras)

    docs = [
        LCDocument(page_content=c, metadata={"source": os.path.basename(DOCX_PATH), "chunk": i})
        for i, c in enumerate(chunks)
    ]

    embeddings = HuggingFaceEmbeddings(model_name=EMBED_MODEL)

    os.makedirs(PERSIST_DIR, exist_ok=True)

    db = Chroma(
        collection_name=COLLECTION,
        embedding_function=embeddings,
        persist_directory=PERSIST_DIR,
    )

    try:
        db.delete_collection()
        db = Chroma(
            collection_name=COLLECTION,
            embedding_function=embeddings,
            persist_directory=PERSIST_DIR,
        )
    except Exception:
        pass

    db.add_documents(docs)

    print(f"✅ Ingested {len(docs)} chunks into '{COLLECTION}' at '{PERSIST_DIR}'")


if __name__ == "__main__":
    main()
